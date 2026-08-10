#!/usr/bin/env python3
"""
customer-bug-report.py — turn a CI run's test failures into a CUSTOMER-READY
bug report: one formal bug per failed test case, with real reproduction steps
and embedded screenshots, delivered as a single PDF (and DOCX).

Every bug follows the agreed customer template:

    Title           [Module] Short description of issue
    Environment     Environment / Platform / Browser / URL / Run
    Preconditions   logged in, module access, fixtures
    Steps to Reproduce   1..N  (the REAL steps the test executed, from the
                                ExtentSpark detailed report when available)
    Actual Result   assertion / exception evidence
    Expected Result derived from the assertion or the test's @Test description
    Severity        heuristic (High / Medium / Low) — conservative by default
    Priority        mirrors severity
    Attachments     screenshots embedded inline (failure screenshot always,
                    plus the last step screenshots before the failure)

Evidence sources (best available wins, per failed test):
  1. ExtentSpark detailed report HTML (reports/detail-report/Detailed_Report_*.html)
     -> real executed steps + inline base64 screenshots + stack trace.
  2. testng-results.xml (reports/groups/**) -> exception message, full stack,
     @Test description, params, timing. Retry-aware: later files override earlier
     ones for the same (class, method, params) — identical semantics to
     consolidated-detailed-report.py.
  3. test-output/screenshots/<method>_FAIL_*.png — matched by method name and
     the invocation's time window (for tests that never reached the Extent report).

Usage (CI summary job):
  python3 .github/scripts/customer-bug-report.py all-reports out-dir \
      --title "Customer Bug Report — Parallel Full Suite" \
      --run-number 123 --run-url https://github.com/... \
      --label "before re-run"

Usage (CI after-rerun job — only STILL-FAILING tests become bugs; recovered
tests are listed in a "recovered on re-run" appendix instead of being filed):
  python3 .github/scripts/customer-bug-report.py all-reports out-dir \
      --rerun-results rerun-results --rerun-detail reports/detail-report \
      --label "after re-run"

Outputs in <out-dir>:
  Customer_Bug_Report.pdf    the deliverable
  Customer_Bug_Report.docx   editable copy (skipped gracefully if python-docx missing)
  Customer_Bug_Report.json   machine-readable bug list (for tooling; never
                             auto-pushed to Jira — per project rule)

Dependencies: reportlab + pillow (hard, for the PDF), python-docx (soft).
"""

import argparse
import base64
import glob
import html as htmllib
import io
import json
import os
import re
import sys
from datetime import datetime, timedelta

# Prefer defusedxml (immune to XXE / billion-laughs) when installed — the CI job
# pip-installs it. The stdlib fallback is additionally guarded by _safe_parse_root
# (refuses any DTD/ENTITY document), matching the repo's other report scripts.
try:
    import defusedxml.ElementTree as ET  # type: ignore
    from xml.etree.ElementTree import ParseError as _XMLParseError
except ImportError:
    import xml.etree.ElementTree as ET
    _XMLParseError = ET.ParseError

# ─────────────────────────────────────────────────────────────────────────
# TestNG results parsing (retry-aware, invocation-level)
# ─────────────────────────────────────────────────────────────────────────

def _safe_parse_root(path):
    """Refuse DTD/ENTITY docs (XXE / billion-laughs need a DOCTYPE; TestNG never
    emits one) — same dependency-free guard as consolidated-detailed-report.py."""
    with open(path, "rb") as fh:
        data = fh.read(50 * 1024 * 1024)
    if b"<!DOCTYPE" in data[:4096] or b"<!ENTITY" in data:
        print(f"  WARNING: skipping {path} — DTD/ENTITY not allowed")
        return None
    return ET.fromstring(data)


def _parse_ts(s):
    """TestNG stamps: '2026-08-08T01:48:50 UTC' or ISO with offset."""
    if not s:
        return None
    s = s.strip().replace(" UTC", "+0000").replace("Z", "+0000")
    for fmt in ("%Y-%m-%dT%H:%M:%S%z", "%Y-%m-%dT%H:%M:%S"):
        try:
            return datetime.strptime(s, fmt).replace(tzinfo=None)
        except ValueError:
            continue
    return None


def parse_invocations(results_dir):
    """{(fqcn, method, params): record} for every non-config invocation.
    Later files override earlier ones for the same key (retry semantics).
    Config-method failures (e.g. classSetup) are collected separately."""
    inv, config_fails = {}, {}
    if not results_dir or not os.path.isdir(results_dir):
        return inv, config_fails
    for path in sorted(glob.glob(os.path.join(results_dir, "**", "testng-results.xml"),
                                 recursive=True)):
        try:
            root = _safe_parse_root(path)
        except (_XMLParseError, ValueError, OSError):
            continue
        if root is None:
            continue
        for cls in root.findall(".//class"):
            fqcn = cls.get("name", "")
            for m in cls.findall("test-method"):
                params = tuple((v.text or "").strip()
                               for v in m.findall("./params/param/value"))
                exc = m.find("exception")
                msg_el = exc.find("message") if exc is not None else None
                stack_el = exc.find("full-stacktrace") if exc is not None else None
                rec = {
                    "fqcn": fqcn,
                    "method": m.get("name", ""),
                    "params": params,
                    "status": m.get("status", "UNKNOWN"),
                    "description": (m.get("description") or "").strip(),
                    "exc_class": exc.get("class", "") if exc is not None else "",
                    "exc_msg": (msg_el.text or "").strip() if msg_el is not None else "",
                    "stack": (stack_el.text or "").strip() if stack_el is not None else "",
                    "started": _parse_ts(m.get("started-at")),
                    "finished": _parse_ts(m.get("finished-at")),
                    "duration_ms": int(m.get("duration-ms") or 0),
                }
                key = (fqcn, rec["method"], params)
                if m.get("is-config") == "true":
                    # keep only FAILED configs; a later PASS of the same config clears it
                    if rec["status"] == "FAIL":
                        config_fails[key] = rec
                    elif key in config_fails and rec["status"] == "PASS":
                        del config_fails[key]
                    continue
                inv[key] = rec
    return inv, config_fails


# ─────────────────────────────────────────────────────────────────────────
# ExtentSpark detailed-report HTML parsing
# ─────────────────────────────────────────────────────────────────────────

NAME_RE = re.compile(r"Detailed_Report_(.+?)_(\d{8}_\d{6})\.html$")
ROW_RE = re.compile(r'<tr class="event-row">(.*?)</tr>', re.S)
TD_RE = re.compile(r"<td[^>]*>(.*?)</td>", re.S)
BADGE_RE = re.compile(r'class="badge log (\w+)-bg"')
IMG_RE = re.compile(r"data:image/(jpeg|jpg|png);base64,([A-Za-z0-9+/=]+)")
TEXTAREA_RE = re.compile(r'<textarea[^>]*class="code-block"[^>]*>(.*?)</textarea>', re.S)
NAME_P_RE = re.compile(r'<p class="name">(.*?)</p>', re.S)
STATUS_ATTR_RE = re.compile(r'status="(\w+)"')
TAG_ATTR_RE = re.compile(r'tag="([^"]*)"')
FAILED_CAPTION_RE = re.compile(r"Test failed: (\w+)")
AT_FRAME_RE = re.compile(r"at (com\.egalvanic[\w.]+)\.(\w+)\(")


def _strip_html(fragment):
    """Visible text of an HTML fragment: drop textareas/imgs, unescape, collapse."""
    frag = TEXTAREA_RE.sub(" ", fragment)
    frag = re.sub(r"<img[^>]*>", " ", frag)
    frag = re.sub(r"<br\s*/?>", "\n", frag)
    frag = re.sub(r"<[^>]+>", " ", frag)
    frag = htmllib.unescape(frag)
    return re.sub(r"[ \t\r\f\v]+", " ", frag).strip()


def parse_extent_items(input_dir, group_hint_from_path=True):
    """All FAIL-badged test items across every Detailed_Report_*.html under
    input_dir. Each item: {group, module, name, tag, rows, methods, raw}.
    rows: [{status, time, text, images[(fmt,b64)], stack}] in execution order."""
    items = []
    for path in sorted(glob.glob(os.path.join(input_dir, "**", "Detailed_Report_*.html"),
                                 recursive=True)):
        m = NAME_RE.search(os.path.basename(path))
        if not m:
            continue
        try:
            if os.path.getsize(path) < 2048:
                continue
            with open(path, encoding="utf-8", errors="replace") as fh:
                doc = fh.read()
        except OSError:
            continue
        module = m.group(1).replace("_", " ").strip()
        rel = os.path.relpath(path, input_dir).replace("\\", "/")
        top = rel.split("/")[0] if "/" in rel else ""
        group = re.sub(r"^reports-(s\d+-)?", "", top) if group_hint_from_path else ""

        for seg in doc.split('<li class="test-item"')[1:]:
            st = STATUS_ATTR_RE.search(seg[:400])
            if not st or st.group(1) != "fail":
                continue
            nm = NAME_P_RE.search(seg)
            name = _strip_html(nm.group(1)) if nm else "(unnamed test)"
            tg = TAG_ATTR_RE.search(seg[:600])
            rows = []
            for row_html in ROW_RE.findall(seg):
                tds = TD_RE.findall(row_html)
                if len(tds) < 3:
                    continue
                badge = BADGE_RE.search(tds[0])
                stack_m = TEXTAREA_RE.search(tds[2])
                rows.append({
                    "status": badge.group(1) if badge else "info",
                    "time": _strip_html(tds[1]),
                    "text": _strip_html(tds[2]),
                    "images": IMG_RE.findall(tds[2]),
                    "stack": htmllib.unescape(stack_m.group(1)).strip() if stack_m else "",
                })
            methods = set(FAILED_CAPTION_RE.findall(seg))
            for fq, meth in AT_FRAME_RE.findall(seg):
                methods.add(meth)
            items.append({"group": group, "module": module, "name": name,
                          "tag": tg.group(1) if tg else "", "rows": rows,
                          "methods": methods, "used": False, "path": path})
    return items


def match_extent_item(items, rec, prefer_dirs=None):
    """Best Extent fail-item for a failed invocation. Match order:
    exact 'Test failed: <method>' / stack frame, then display-name heuristic.
    Data-driven params disambiguate via the display name. prefer_dirs: item
    paths starting with any of these are tried first (fresh re-run reports)."""
    method = rec["method"]
    stripped = method[4:] if method.startswith("test") else method
    shown_params = [p for p in rec["params"]
                    if p and not re.match(PARAM_NOISE_PATTERN, p)]

    def tiers(pool):
        exact = [it for it in pool if method in it["methods"]]
        byname = [it for it in pool
                  if stripped and stripped in it["name"] and it not in exact]
        return exact + byname

    def pick(pool):
        cands = tiers([it for it in pool if not it["used"]])
        if not cands:
            return None
        if shown_params:
            for it in cands:
                if all(p in it["name"] for p in shown_params):
                    return it
        return cands[0]

    if prefer_dirs:
        preferred = [it for it in items
                     if any(it["path"].startswith(d) for d in prefer_dirs)]
        hit = pick(preferred)
        if hit:
            return hit
    return pick(items)


# ─────────────────────────────────────────────────────────────────────────
# Fallback screenshots (test-output/screenshots/<method>_FAIL_<ts>.png)
# ─────────────────────────────────────────────────────────────────────────

SHOT_TS_RE = re.compile(r"_(\d{8}_\d{6})\.(png|jpg|jpeg)$", re.I)


def index_screenshots(input_dir, extra_dirs=()):
    """[(basename, path, ts_datetime_or_None)] for every screenshot file under
    input_dir's screenshots/ folders, plus any extra dirs scanned wholesale
    (e.g. the re-run job's local test-output/screenshots)."""
    paths = list(glob.glob(os.path.join(input_dir, "**", "screenshots", "*.*"),
                           recursive=True))
    for d in extra_dirs or ():
        if d and os.path.isdir(d):
            paths.extend(glob.glob(os.path.join(d, "**", "*.*"), recursive=True))
    out, seen = [], set()
    for path in paths:
        base = os.path.basename(path)
        if path in seen or not re.search(r"\.(png|jpe?g)$", base, re.I):
            continue
        seen.add(path)
        m = SHOT_TS_RE.search(base)
        ts = None
        if m:
            try:
                ts = datetime.strptime(m.group(1), "%Y%m%d_%H%M%S")
            except ValueError:
                pass
        out.append((base, path, ts))
    return out


def fallback_screenshots(shots, rec, limit):
    """Screenshots whose filename starts with the failed method (or method minus
    'test') — FAIL-tagged first, then constrained to the invocation's time window
    when both sides carry timestamps."""
    method = rec["method"]
    stripped = method[4:] if method.startswith("test") else method
    cands = []
    for base, path, ts in shots:
        stem = base.rsplit(".", 1)[0]
        if not (stem.startswith(method) or (stripped and stem.startswith(stripped))):
            continue
        in_window = True
        if ts and rec["started"] and rec["finished"]:
            in_window = (rec["started"] - timedelta(seconds=10)
                         <= ts <= rec["finished"] + timedelta(seconds=180))
        is_fail = "_FAIL" in base.upper()
        cands.append((not is_fail, not in_window, ts or datetime.min, path))
    cands.sort()
    return [c[3] for c in cands[:limit]]


# ─────────────────────────────────────────────────────────────────────────
# Bug model
# ─────────────────────────────────────────────────────────────────────────

EXPECTED_SPLIT_RE = re.compile(r"^(.*?)\s*expected\s*\[(.*?)\]\s*but found\s*\[(.*?)\]",
                               re.S | re.I)

# Java object/array toString refs are noise, never customer data:
#   com.foo.Page@1a2b  ·  [Ljava.lang.String;@a22c4d8  ·  [[I@5e91993f
PARAM_NOISE_PATTERN = r"^\[*[\w.$;]+@[0-9a-fA-F]+$"

# Java/Selenium internals that mean nothing to a customer.
LAMBDA_NOISE_RE = re.compile(r"[\w.$]+\$\$Lambda\$\d+/0x[0-9a-fA-F]+@[0-9a-fA-F]+")
OBJREF_NOISE_RE = re.compile(r"\b[\w.]+(?:Page|Util|Helper)@[0-9a-fA-F]+\b")
BUILDINFO_RE = re.compile(r"Build info: .*$|Driver info: .*$|System info: .*$",
                          re.S | re.M)


def sanitize_customer_text(s):
    """Strip Java lambda refs / Selenium build-info blocks from customer-visible text."""
    s = LAMBDA_NOISE_RE.sub("the expected page state", s or "")
    s = OBJREF_NOISE_RE.sub("the page", s)
    s = BUILDINFO_RE.sub("", s)
    return re.sub(r"[ \t]+", " ", s).strip()


def tc_id_of(rec):
    """Human test-case id from the method / description, e.g. TC_WOP_025, ENG_07,
    AS08 — used to keep titles unique across data-driven variants."""
    for src in (rec["method"], rec["description"]):
        m = re.search(r"(TC[_-][A-Za-z0-9_]+?|[A-Z]{2,6}[_-]?\d{1,4}[a-z]?)(?=_|:|\s|$)",
                      src or "")
        if m:
            return m.group(1)
    return rec["method"]


def derive_module(item, rec):
    if item:
        return item["module"]
    short = rec["fqcn"].rsplit(".", 1)[-1]
    short = re.sub(r"TestNG$|Test$", "", short)
    words = re.sub(r"(?<!^)(?=[A-Z][a-z])", " ", short).strip()
    # collapse suite-part suffixes like "Asset Part1" -> "Asset"
    words = re.sub(r"\s*Part\s*\d.*$", "", words)
    return words or short


def first_line(s, limit=160):
    line = (s or "").strip().splitlines()[0] if (s or "").strip() else ""
    return (line[: limit - 1] + "…") if len(line) > limit else line


def derive_short_issue(rec):
    """Human phrase for the Title, from the strongest available signal."""
    msg = sanitize_customer_text(rec["exc_msg"])
    m = EXPECTED_SPLIT_RE.match(msg or "")
    if m and m.group(1).strip():
        return first_line(m.group(1).strip().rstrip(":"), 90)
    cm = re.match(r"Browser console reported (\d+) severe error", msg or "")
    if cm:
        ctx = re.search(r"during \[([^\]]+)\]", msg)
        return f"Console shows {cm.group(1)} severe error(s)" + (
            f" on {ctx.group(1)}" if ctx else "")
    if "TimeoutException" in (rec["exc_class"] or ""):
        # the raw wait-condition text is Java noise — the @Test description says
        # what the user-visible expectation was
        if rec["description"]:
            return "Page did not load in time — " + first_line(rec["description"], 70)
        return "Page/element did not load within the expected time"
    if msg:
        return first_line(msg, 90)
    if rec["description"]:
        return first_line(rec["description"], 90) + " — failed"
    return rec["method"] + " failed"


def derive_expected(rec):
    msg = sanitize_customer_text(rec["exc_msg"]) or ""
    m = EXPECTED_SPLIT_RE.match(msg)
    if m:
        head = m.group(1).strip().rstrip(":")
        exp, act = m.group(2).strip(), m.group(3).strip()
        if head:
            return (f"{head}. The check should evaluate to [{exp}], "
                    f"but the application produced [{act}].")
        return f"The check should evaluate to [{exp}], but the application produced [{act}]."
    if rec["description"]:
        return (f"The scenario “{rec['description']}” should complete successfully "
                f"without errors.")
    if "TimeoutException" in (rec["exc_class"] or ""):
        return ("The expected element/condition should appear within the normal "
                "loading time; the page should not hang or stay incomplete.")
    return (f"The flow should complete without raising "
            f"{rec['exc_class'].rsplit('.', 1)[-1] or 'an error'}.")


# 5xx only in an HTTP-status context — a bare number match would also hit
# noise like "tried for 20 second(s) with 500 milliseconds interval".
HIGH_SIGNALS = re.compile(
    r"status (?:of )?5\d{2}\b|HTTP[ /]5\d{2}\b|\b5\d{2} (?:Internal|Bad Gateway|"
    r"Service Unavailable)|Internal Server Error|psycopg2|SQLSTATE|"
    r"Application Error|OutOfMemory", re.I)
CONSOLE_SEVERE_RE = re.compile(r"Browser console reported (\d+) severe")


def derive_severity(rec, is_config=False, reproduced_on_rerun=False):
    blob = " ".join([rec["exc_msg"] or "", (rec["stack"] or "")[:2000]])
    if is_config:
        return "High", ("Test-class setup failed — every test in the class was "
                        "blocked. Usually an environment/login/navigation breakage.")
    if HIGH_SIGNALS.search(blob):
        return "High", "Server-side error signals (5xx / backend exception) in the evidence."
    m = CONSOLE_SEVERE_RE.search(rec["exc_msg"] or "")
    if m and int(m.group(1)) >= 3:
        return "Medium", (f"{m.group(1)} severe browser-console errors captured — "
                          "API/auth errors visible to the end user's session.")
    if "TimeoutException" in (rec["exc_class"] or "") or \
       "NoSuchElement" in (rec["exc_class"] or "") or \
       "StaleElement" in (rec["exc_class"] or ""):
        note = "Element/timing failure — verify manually; can also indicate slowness."
        if reproduced_on_rerun:
            note = "Element/timing failure that REPRODUCED on a clean re-run."
        return "Medium", note
    return "Medium", ("Functional assertion failed."
                      + (" Reproduced on re-run (not flaky)." if reproduced_on_rerun else ""))


HEURISTIC_ACTION_STEPS = [
    (re.compile(r"create|add", re.I), ["Open the Create / Add form for the relevant entity.",
                                       "Fill the required fields and submit."]),
    (re.compile(r"delete|remove", re.I), ["Locate the target record in the list.",
                                          "Trigger the Delete action and confirm."]),
    (re.compile(r"edit|update", re.I), ["Open the existing record's Edit drawer/form.",
                                        "Modify the relevant field and save."]),
    (re.compile(r"login|auth|signin", re.I), ["Go to the login page.",
                                              "Enter the credentials matching the test inputs."]),
]


def heuristic_steps(rec, module, base_url):
    steps = [f"Log in to {base_url} with a standard QA user.",
             f"Navigate to the {module} module."]
    name = rec["method"] + " " + (rec["description"] or "")
    for pat, extra in HEURISTIC_ACTION_STEPS:
        if pat.search(name):
            steps.extend(extra)
            break
    else:
        steps.append("Execute the scenario: "
                     + (rec["description"] or rec["method"])
                     + " (see test reference below for the exact automated flow).")
    steps.append("Observe the result described under “Actual Result”.")
    return steps


def shown_params_of(rec):
    return [p for p in rec["params"]
            if p and not re.match(PARAM_NOISE_PATTERN, p)]


def build_bug(rec, item, shots, args, is_config=False, reproduced_on_rerun=False):
    """Assemble one customer-template bug dict from the best available evidence."""
    module = derive_module(item, rec)
    short = derive_short_issue(rec)
    params = shown_params_of(rec)
    title = f"[{module}] {short}"
    if params:
        title += " · " + ", ".join(params[:3])

    # Steps + screenshots
    images = []          # list of (bytes_or_b64, kind, caption); kind: 'b64jpeg'|'file'
    steps = []
    if item and item["rows"]:
        for row in item["rows"]:
            if row["status"] in ("info", "pass", "warning") and row["text"]:
                steps.append(row["text"])
            for fmt, b64 in row["images"]:
                cap = ("Failure screenshot" if row["status"] == "fail"
                       else f"Step {len(steps)}: {first_line(row['text'], 80)}")
                images.append((b64, "b64", cap, row["status"] == "fail"))
        if not steps:
            steps = heuristic_steps(rec, module, args.base_url)
        else:
            steps = [first_line(s, 220) for s in steps]
            steps.append("Observe the failure described under “Actual Result”.")
    else:
        steps = heuristic_steps(rec, module, args.base_url)
        for path in fallback_screenshots(shots, rec, args.max_screenshots):
            images.append((path, "file", "Failure screenshot (captured at test failure)",
                           True))

    # keep the failure screenshot(s) + the last steps before it, capped
    fail_imgs = [im for im in images if im[3]]
    step_imgs = [im for im in images if not im[3]]
    budget = max(args.max_screenshots - len(fail_imgs[:2]), 0)
    images = step_imgs[-budget:] + fail_imgs[:2] if budget else fail_imgs[:2]

    # Actual result
    actual_parts = []
    if rec["exc_msg"]:
        clean_msg = sanitize_customer_text(rec["exc_msg"])
        actual_parts.append(first_line(clean_msg, 600) if "\n" not in clean_msg
                            else clean_msg[:900])
    elif item:
        fail_texts = [r["text"] for r in item["rows"] if r["status"] == "fail" and r["text"]]
        if fail_texts:
            actual_parts.append(fail_texts[0][:600])
    if rec["exc_class"]:
        actual_parts.append(f"(technical: {rec['exc_class'].rsplit('.', 1)[-1]})")
    actual = "\n".join(actual_parts) or "The step failed — see the attached screenshot."

    severity, severity_note = derive_severity(rec, is_config, reproduced_on_rerun)

    preconditions = ["User is logged in with a standard QA account.",
                     f"User has access to the {module} module."]
    if is_config:
        preconditions.append("No further precondition — the failure happens while "
                             "preparing the module itself (login / first navigation).")

    stack = rec["stack"]
    if not stack and item:
        for r in item["rows"]:
            if r["stack"]:
                stack = r["stack"]
                break

    return {
        "title": title,
        "module": module,
        "severity": severity,
        "severity_note": severity_note,
        "priority": severity,
        "preconditions": preconditions,
        "steps": steps,
        "actual": actual,
        "expected": derive_expected(rec),
        "images": images,
        "tc_id": tc_id_of(rec),
        "test_ref": f"{rec['fqcn']}#{rec['method']}"
                    + (f" [{', '.join(params)}]" if params else ""),
        "duration_ms": rec["duration_ms"],
        "reproduced_on_rerun": reproduced_on_rerun,
        "is_config": is_config,
        "stack_head": "\n".join((stack or "").splitlines()[:12]),
        "evidence": "extent" if item else "testng+screenshot",
    }


# ─────────────────────────────────────────────────────────────────────────
# Image normalization (everything → JPEG bytes, bounded width)
# ─────────────────────────────────────────────────────────────────────────

def normalize_image(spec, max_width, quality):
    """(payload, kind) → JPEG bytes + (w, h), or None. Uses Pillow."""
    from PIL import Image
    payload, kind = spec
    try:
        raw = base64.b64decode(payload) if kind == "b64" else open(payload, "rb").read()
        img = Image.open(io.BytesIO(raw))
        img.load()
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        if img.width > max_width:
            ratio = max_width / img.width
            img = img.resize((max_width, max(1, int(img.height * ratio))))
        buf = io.BytesIO()
        img.save(buf, "JPEG", quality=quality, optimize=True)
        return buf.getvalue(), img.size
    except Exception as e:
        print(f"  WARNING: could not process an attachment image: {e}")
        return None


# ─────────────────────────────────────────────────────────────────────────
# PDF rendering (reportlab / platypus)
# ─────────────────────────────────────────────────────────────────────────

SEV_COLORS = {"High": "#c0392b", "Medium": "#b9770e", "Low": "#1e8449"}


def render_pdf(bugs, recovered, args, out_path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (BaseDocTemplate, Frame, Image, PageBreak,
                                    PageTemplate, Paragraph, Spacer, Table,
                                    TableStyle, KeepTogether)
    from reportlab.lib.utils import ImageReader

    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=9.5, leading=13)
    small = ParagraphStyle("small", parent=body, fontSize=8, textColor=colors.grey)
    h1 = ParagraphStyle("h1x", parent=styles["Heading1"], fontSize=20, spaceAfter=6)
    h2 = ParagraphStyle("h2x", parent=styles["Heading2"], fontSize=13,
                        spaceBefore=10, spaceAfter=4)
    lbl = ParagraphStyle("lbl", parent=body, fontName="Helvetica-Bold")
    mono = ParagraphStyle("mono", parent=body, fontName="Courier", fontSize=7.5,
                          leading=9.5, backColor=colors.HexColor("#f4f4f4"))

    def esc(s):
        return htmllib.escape(s or "").replace("\n", "<br/>")

    page_w, page_h = A4
    margin = 15 * mm
    frame_w = page_w - 2 * margin

    doc = BaseDocTemplate(out_path, pagesize=A4,
                          leftMargin=margin, rightMargin=margin,
                          topMargin=margin, bottomMargin=margin + 4 * mm,
                          title=args.title, author="eGalvanic QA Automation")

    def on_page(canv, _doc):
        canv.saveState()
        canv.setFont("Helvetica", 7)
        canv.setFillColor(colors.grey)
        canv.drawString(margin, 8 * mm, args.title)
        canv.drawRightString(page_w - margin, 8 * mm, f"Page {canv.getPageNumber()}")
        canv.restoreState()

    doc.addPageTemplates([PageTemplate(
        id="all", frames=[Frame(margin, margin + 4 * mm, frame_w,
                                page_h - 2 * margin - 4 * mm)],
        onPage=on_page)])

    story = []

    # ── Cover / summary
    sev_counts = {}
    for b in bugs:
        sev_counts[b["severity"]] = sev_counts.get(b["severity"], 0) + 1
    story.append(Paragraph(esc(args.title), h1))
    meta_rows = [
        ["Report generated", args.run_date],
        ["Scope", args.label or "full run"],
        ["Environment", f"{args.environment} — {args.base_url}"],
        ["Platform / Browser", f"Web — {args.browser}"],
        ["CI run", f"#{args.run_number}  {args.run_url}".strip()],
        ["Bugs in this report", f"{len(bugs)}  ("
         + ", ".join(f"{v} {k}" for k, v in sorted(sev_counts.items())) + ")"
         if bugs else "0 — all tests passed"],
    ]
    if recovered:
        meta_rows.append(["Recovered on re-run (not filed)", str(len(recovered))])
    t = Table([[Paragraph(esc(k), lbl), Paragraph(esc(v), body)] for k, v in meta_rows],
              colWidths=[45 * mm, frame_w - 45 * mm])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f3f7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(t)
    story.append(Spacer(1, 6 * mm))

    if bugs:
        story.append(Paragraph("Summary of reported bugs", h2))
        rows = [[Paragraph("<b>#</b>", body), Paragraph("<b>Title</b>", body),
                 Paragraph("<b>Severity</b>", body)]]
        for i, b in enumerate(bugs, 1):
            rows.append([Paragraph(f"BUG-{i:03d}", body),
                         Paragraph(esc(b["title"]), body),
                         Paragraph(f'<font color="{SEV_COLORS[b["severity"]]}">'
                                   f'{b["severity"]}</font>', body)])
        t = Table(rows, colWidths=[21 * mm, frame_w - 21 * mm - 20 * mm, 20 * mm],
                  repeatRows=1)
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2a44")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#f7f9fb")]),
            ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 2.5), ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
        ]))
        story.append(t)

    # ── One bug per section
    for i, b in enumerate(bugs, 1):
        story.append(PageBreak())
        sev_color = SEV_COLORS[b["severity"]]
        head = Table([[Paragraph(f"<b>BUG-{i:03d}</b>", ParagraphStyle(
                          "bug", parent=body, textColor=colors.white, fontSize=11)),
                       Paragraph(f'<b>{esc(b["title"])}</b>', ParagraphStyle(
                          "bugt", parent=body, textColor=colors.white, fontSize=11))]],
                     colWidths=[22 * mm, frame_w - 22 * mm])
        head.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#1f2a44")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(head)
        story.append(Spacer(1, 3 * mm))

        env_rows = [
            ["Environment", f"{args.environment}"],
            ["Platform", "Web"],
            ["Browser/App Version", args.browser + (f" — app {args.app_version}"
                                                    if args.app_version else "")],
            ["URL", args.base_url],
            ["Device", "Desktop (CI: Linux, 1920×1080)"],
            ["Severity", b["severity"] + " — " + b["severity_note"]],
            ["Priority", b["priority"]],
        ]
        if b["reproduced_on_rerun"]:
            env_rows.append(["Reproducibility",
                             "Failed in the parallel run AND on a clean re-run "
                             "(reproducible, not flaky)"])
        t = Table([[Paragraph(esc(k), lbl), Paragraph(esc(v), body)]
                   for k, v in env_rows],
                  colWidths=[42 * mm, frame_w - 42 * mm])
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f3f7")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        # severity cell color
        t.setStyle(TableStyle([("TEXTCOLOR", (1, 5), (1, 5),
                                colors.HexColor(sev_color))]))
        story.append(t)

        story.append(Paragraph("Preconditions", h2))
        for p in b["preconditions"]:
            story.append(Paragraph("• " + esc(p), body))

        story.append(Paragraph("Steps to Reproduce", h2))
        for n, s in enumerate(b["steps"], 1):
            story.append(Paragraph(f"{n}. {esc(s)}", body))

        story.append(Paragraph("Actual Result", h2))
        story.append(Paragraph(esc(b["actual"]), body))

        story.append(Paragraph("Expected Result", h2))
        story.append(Paragraph(esc(b["expected"]), body))

        if b["stack_head"] and args.include_stack:
            story.append(Paragraph("Technical evidence (for the dev team)", h2))
            story.append(Paragraph(esc(b["stack_head"]), mono))

        if b["images"]:
            story.append(Paragraph("Attachments", h2))
            for (payload, kind, caption, _isfail) in b["images"]:
                norm = normalize_image((payload, kind), args.max_image_width,
                                       args.jpeg_quality)
                if not norm:
                    continue
                data, (w, h) = norm
                disp_w = min(frame_w, w)
                disp_h = h * (disp_w / w)
                max_h = page_h - 2 * margin - 30 * mm
                if disp_h > max_h:
                    disp_w *= max_h / disp_h
                    disp_h = max_h
                story.append(KeepTogether([
                    Image(io.BytesIO(data), width=disp_w, height=disp_h),
                    Paragraph(esc(caption), small),
                ]))
                story.append(Spacer(1, 2 * mm))

        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(
            f"Test reference: {esc(b['test_ref'])} · evidence: {b['evidence']}"
            + (" · test-class setup failure" if b["is_config"] else ""), small))

    # ── Appendix: recovered-on-rerun (flaky) tests
    if recovered:
        story.append(PageBreak())
        story.append(Paragraph("Appendix — failed once but recovered on re-run "
                               "(not filed as bugs)", h2))
        story.append(Paragraph(
            "These test cases failed during the parallel run but PASSED when "
            "re-run in isolation. They are treated as flaky/environmental, not "
            "as product bugs.", body))
        story.append(Spacer(1, 2 * mm))
        rows = [[Paragraph("<b>Test</b>", body)]]
        for k in recovered:
            rows.append([Paragraph(esc(k), body)])
        t = Table(rows, colWidths=[frame_w], repeatRows=1)
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2a44")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 2.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
        ]))
        story.append(t)

    doc.build(story)


# ─────────────────────────────────────────────────────────────────────────
# DOCX rendering (best-effort)
# ─────────────────────────────────────────────────────────────────────────

def render_docx(bugs, recovered, args, out_path):
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        print("[customer-bug-report] python-docx not installed — skipping DOCX "
              "(PDF is still generated)")
        return False

    SEV_RGB = {"High": RGBColor(0xC0, 0x39, 0x2B),
               "Medium": RGBColor(0xB9, 0x77, 0x0E),
               "Low": RGBColor(0x1E, 0x84, 0x49)}
    d = docx.Document()
    d.add_heading(args.title, level=0)
    p = d.add_paragraph()
    p.add_run(f"Generated {args.run_date} · {args.label or 'full run'} · "
              f"{args.environment} — {args.base_url} · CI run #{args.run_number}\n"
              f"{args.run_url}").font.size = Pt(9)
    d.add_paragraph(f"Bugs in this report: {len(bugs)}"
                    + (f" · recovered on re-run (not filed): {len(recovered)}"
                       if recovered else ""))

    for i, b in enumerate(bugs, 1):
        d.add_page_break()
        d.add_heading(f"BUG-{i:03d}  {b['title']}", level=1)
        table = d.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        fields = [("Environment", args.environment), ("Platform", "Web"),
                  ("Browser/App Version", args.browser
                   + (f" — app {args.app_version}" if args.app_version else "")),
                  ("URL", args.base_url),
                  ("Device", "Desktop (CI: Linux, 1920×1080)"),
                  ("Severity", f"{b['severity']} — {b['severity_note']}"),
                  ("Priority", b["priority"])]
        if b["reproduced_on_rerun"]:
            fields.append(("Reproducibility", "Failed in the parallel run AND on a "
                           "clean re-run (reproducible, not flaky)"))
        for k, v in fields:
            row = table.add_row()
            row.cells[0].text = k
            row.cells[1].text = v
            if k == "Severity":
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.color.rgb = SEV_RGB[b["severity"]]
        d.add_heading("Preconditions", level=2)
        for pre in b["preconditions"]:
            d.add_paragraph(pre, style="List Bullet")
        d.add_heading("Steps to Reproduce", level=2)
        for s in b["steps"]:
            d.add_paragraph(s, style="List Number")
        d.add_heading("Actual Result", level=2)
        d.add_paragraph(b["actual"])
        d.add_heading("Expected Result", level=2)
        d.add_paragraph(b["expected"])
        if b["stack_head"] and args.include_stack:
            d.add_heading("Technical evidence (for the dev team)", level=2)
            pr = d.add_paragraph(b["stack_head"])
            for run in pr.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(7.5)
        if b["images"]:
            d.add_heading("Attachments", level=2)
            for (payload, kind, caption, _isfail) in b["images"]:
                norm = normalize_image((payload, kind), args.max_image_width,
                                       args.jpeg_quality)
                if not norm:
                    continue
                data, _size = norm
                try:
                    d.add_picture(io.BytesIO(data), width=Inches(6.2))
                    cap = d.add_paragraph(caption)
                    cap.runs[0].font.size = Pt(8)
                except Exception as e:
                    print(f"  WARNING: docx image skipped: {e}")
        d.add_paragraph(f"Test reference: {b['test_ref']}").runs[0].font.size = Pt(8)

    if recovered:
        d.add_page_break()
        d.add_heading("Appendix — failed once but recovered on re-run "
                      "(not filed as bugs)", level=1)
        for k in recovered:
            d.add_paragraph(k, style="List Bullet")

    d.save(out_path)
    return True


# ─────────────────────────────────────────────────────────────────────────
# Orchestration
# ─────────────────────────────────────────────────────────────────────────

def short_key(key):
    fqcn, method, params = key
    label = f"{fqcn.rsplit('.', 1)[-1]}.{method}"
    shown = [p for p in params if p and not re.match(PARAM_NOISE_PATTERN, p)]
    if shown:
        label += "(" + ", ".join(shown[:3]) + ("…" if len(shown) > 3 else "") + ")"
    return label


def main():
    ap = argparse.ArgumentParser(description="Customer-facing per-run bug report "
                                             "(PDF + DOCX) from CI test artifacts.")
    ap.add_argument("input_dir", help="dir with the run's artifacts (CI: all-reports)")
    ap.add_argument("output_dir")
    ap.add_argument("--title", default="Customer Bug Report — eGalvanic Web QA")
    ap.add_argument("--label", default="", help="e.g. 'before re-run' / 'after re-run'")
    ap.add_argument("--run-number", default="", dest="run_number")
    ap.add_argument("--run-url", default="", dest="run_url")
    ap.add_argument("--run-date", default="", dest="run_date")
    ap.add_argument("--environment", default="QA")
    ap.add_argument("--base-url", default="https://acme.qa.egalvanic.ai",
                    dest="base_url")
    ap.add_argument("--browser", default="Google Chrome (latest stable, "
                                         "GitHub Actions Linux runner)")
    ap.add_argument("--app-version", default="", dest="app_version")
    ap.add_argument("--rerun-results", default=None,
                    help="dir with the re-run's testng-results.xml. When given, only "
                         "STILL-FAILING tests become bugs; recovered → appendix.")
    ap.add_argument("--rerun-detail", default=None,
                    help="dir with the re-run's Detailed_Report_*.html (preferred "
                         "step/screenshot source for still-failing tests)")
    ap.add_argument("--extra-screenshots", action="append", default=[],
                    help="additional screenshot dir(s) to index (repeatable), "
                         "e.g. the re-run job's local test-output/screenshots")
    ap.add_argument("--max-screenshots", type=int, default=4)
    ap.add_argument("--max-image-width", type=int, default=1300)
    ap.add_argument("--jpeg-quality", type=int, default=70)
    ap.add_argument("--max-bugs", type=int, default=0,
                    help="safety cap; 0 = unlimited")
    ap.add_argument("--include-stack", action="store_true", default=True)
    ap.add_argument("--no-stack", dest="include_stack", action="store_false")
    ap.add_argument("--formats", default="pdf,docx")
    args = ap.parse_args()
    if not args.run_date:
        args.run_date = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    # 1. what failed?
    inv, config_fails = parse_invocations(args.input_dir)
    recovered = []
    if args.rerun_results:
        rerun_inv, rerun_cfg = parse_invocations(args.rerun_results)
        orig_failed = {k for k, r in inv.items() if r["status"] == "FAIL"}
        merged = dict(inv)
        for k, r in rerun_inv.items():
            merged[k] = r
        failed_keys = sorted(k for k, r in merged.items() if r["status"] == "FAIL")
        recovered = sorted(short_key(k) for k in orig_failed
                           if rerun_inv.get(k, {}).get("status") == "PASS")
        records = {k: merged[k] for k in failed_keys}
        # a config failure counts as cleared if the re-run has any result for its class
        rerun_classes = {k[0] for k in rerun_inv}
        config_fails = {k: v for k, v in config_fails.items()
                        if k[0] not in rerun_classes or k in rerun_cfg}
        config_fails.update(rerun_cfg)
    else:
        records = {k: r for k, r in inv.items() if r["status"] == "FAIL"}

    print(f"[customer-bug-report] failed test invocations: {len(records)} "
          f"| failed configs: {len(config_fails)} | recovered on re-run: {len(recovered)}")

    # 2. evidence
    items = parse_extent_items(args.input_dir)
    prefer = None
    if args.rerun_detail and os.path.isdir(args.rerun_detail):
        items += parse_extent_items(args.rerun_detail, group_hint_from_path=False)
        prefer = [os.path.abspath(args.rerun_detail), args.rerun_detail]
    shots = index_screenshots(args.input_dir, args.extra_screenshots)
    print(f"[customer-bug-report] extent fail-items: {len(items)} "
          f"| screenshots indexed: {len(shots)}")

    # 3. build bugs (stable order: module, then test ref)
    bugs = []
    for key in sorted(records, key=lambda k: (k[0], k[1], k[2])):
        rec = records[key]
        try:
            item = match_extent_item(items, rec, prefer_dirs=prefer)
            if item:
                item["used"] = True
            bugs.append(build_bug(rec, item, shots, args,
                                  reproduced_on_rerun=bool(args.rerun_results)))
        except Exception as e:
            print(f"  WARNING: bug build failed for {short_key(key)}: {e}")
    for key in sorted(config_fails, key=lambda k: (k[0], k[1])):
        rec = config_fails[key]
        try:
            bugs.append(build_bug(rec, None, shots, args, is_config=True,
                                  reproduced_on_rerun=bool(args.rerun_results)))
        except Exception as e:
            print(f"  WARNING: config-bug build failed for {short_key(key)}: {e}")

    bugs.sort(key=lambda b: ({"High": 0, "Medium": 1, "Low": 2}[b["severity"]],
                             b["module"], b["title"]))

    # data-driven variants can produce identical titles — suffix the test-case id
    # so every bug in the customer report is uniquely identifiable
    title_counts = {}
    for b in bugs:
        title_counts[b["title"]] = title_counts.get(b["title"], 0) + 1
    for b in bugs:
        if title_counts[b["title"]] > 1 and b["tc_id"] not in b["title"]:
            b["title"] += f" ({b['tc_id']})"

    if args.max_bugs and len(bugs) > args.max_bugs:
        print(f"[customer-bug-report] capping {len(bugs)} bugs at --max-bugs="
              f"{args.max_bugs}")
        bugs = bugs[:args.max_bugs]

    os.makedirs(args.output_dir, exist_ok=True)

    # 4. machine-readable copy (images stripped)
    json_path = os.path.join(args.output_dir, "Customer_Bug_Report.json")
    with open(json_path, "w", encoding="utf-8") as fh:
        json.dump({"title": args.title, "label": args.label,
                   "run_number": args.run_number, "run_url": args.run_url,
                   "generated": args.run_date, "recovered_on_rerun": recovered,
                   "bugs": [{k: v for k, v in b.items() if k != "images"}
                            | {"attachment_count": len(b["images"])}
                            for b in bugs]},
                  fh, indent=2, ensure_ascii=False)
    print(f"[customer-bug-report] wrote {json_path}")

    formats = {f.strip().lower() for f in args.formats.split(",") if f.strip()}
    if "pdf" in formats:
        pdf_path = os.path.join(args.output_dir, "Customer_Bug_Report.pdf")
        render_pdf(bugs, recovered, args, pdf_path)
        print(f"[customer-bug-report] wrote {pdf_path} "
              f"({os.path.getsize(pdf_path) / 1024 / 1024:.1f} MB, {len(bugs)} bugs)")
    if "docx" in formats:
        docx_path = os.path.join(args.output_dir, "Customer_Bug_Report.docx")
        if render_docx(bugs, recovered, args, docx_path):
            print(f"[customer-bug-report] wrote {docx_path} "
                  f"({os.path.getsize(docx_path) / 1024 / 1024:.1f} MB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
